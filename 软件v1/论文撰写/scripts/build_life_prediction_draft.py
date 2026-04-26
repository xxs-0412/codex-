from __future__ import annotations

import math
import os
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


BASE_DIR = Path(__file__).resolve().parents[2]
WORK_DIR = BASE_DIR / "论文撰写"
FIG_DIR = WORK_DIR / "figures"
OUT_DOCX = WORK_DIR / "寿命预测系统架构论文初稿.docx"

RESULT_DIR = BASE_DIR / "结果" / "4.25测试集选定_论文结构重跑"
BEAUTY_DIR = RESULT_DIR / "论文图表_美化版"
DATA_SUMMARY = BASE_DIR / "仿真数据" / "数据集（处理后）" / "修改数据汇总.csv"
SPLIT_CSV = RESULT_DIR / "00_实验配置快照" / "正式测试集划分.csv"

BODY_FONT = "SimSun"
HEAD_FONT = "Microsoft YaHei"
ACCENT = "1F4E79"
ACCENT_LIGHT = "EAF2F8"
SUBTLE = "F5F7FA"
GRID = "D8DEE9"
TABLE_IMG_COUNTER = 0


def font_path(name: str) -> str | None:
    candidates = [
        Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / name,
        Path(r"C:\Windows\Fonts") / name,
    ]
    for p in candidates:
        if p.exists():
            return str(p)
    return None


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    choices = ["msyhbd.ttc" if bold else "msyh.ttc", "simhei.ttf", "simsun.ttc"]
    for name in choices:
        p = font_path(name)
        if p:
            return ImageFont.truetype(p, size)
    return ImageFont.load_default()


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> tuple[int, int]:
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for paragraph in text.split("\n"):
        if not paragraph:
            lines.append("")
            continue
        current = ""
        for ch in paragraph:
            trial = current + ch
            if text_size(draw, trial, font)[0] <= max_width or not current:
                current = trial
            else:
                lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines


def rounded_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    fill: str,
    outline: str,
    title: str,
    body: str = "",
    title_color: str = "#102A43",
) -> None:
    x0, y0, x1, y1 = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    title_font = get_font(24, True)
    body_font = get_font(20)
    tx = x0 + 20
    ty = y0 + 16
    draw.text((tx, ty), title, font=title_font, fill=title_color)
    if body:
        lines = wrap_text(draw, body, body_font, x1 - x0 - 40)
        y = ty + 38
        for line in lines[:6]:
            draw.text((tx, y), line, font=body_font, fill="#334E68")
            y += 28


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#34495E", width: int = 4) -> None:
    draw.line([start, end], fill=color, width=width)
    sx, sy = start
    ex, ey = end
    angle = math.atan2(ey - sy, ex - sx)
    head_len = 18
    spread = math.pi / 7
    points = [
        (ex, ey),
        (ex - head_len * math.cos(angle - spread), ey - head_len * math.sin(angle - spread)),
        (ex - head_len * math.cos(angle + spread), ey - head_len * math.sin(angle + spread)),
    ]
    draw.polygon(points, fill=color)


def make_architecture_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 1080), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(38, True)
    small_font = get_font(20)
    draw.text((70, 42), "寿命预测系统架构：物理增强 Transformer 闭环代理模型", font=title_font, fill="#102A43")
    draw.text((72, 94), "输入工况与磨损状态 -> 应力预测 -> Archard 磨损更新 -> 寿命阈值判定", font=small_font, fill="#52616B")

    # Top traditional model strip
    rounded_box(draw, (80, 155, 430, 300), "#F8F9FB", "#BCCCDC", "传统神经网络模型", "输入 F, D, Cr 等静态工况，直接回归寿命 L_pred")
    rounded_box(draw, (540, 155, 885, 300), "#EFF6FF", "#74A9D8", "FNN / RNN / CNN", "缺少磨损状态递推，物理过程表达较弱")
    rounded_box(draw, (995, 155, 1295, 300), "#FFF7E6", "#E0A458", "寿命输出", "L_pred")
    arrow(draw, (430, 228), (540, 228))
    arrow(draw, (885, 228), (995, 228))

    # Main enhanced pipeline
    draw.rounded_rectangle((55, 365, 1745, 990), radius=28, fill="#F6FAFD", outline="#D6E4F0", width=3)
    draw.text((90, 392), "本文采用的物理增强代理模型", font=get_font(30, True), fill="#1F4E79")

    boxes = {
        "input": (95, 480, 345, 650),
        "feature": (420, 480, 700, 650),
        "stress": (780, 480, 1080, 650),
        "archard": (1150, 480, 1410, 650),
        "state": (1480, 480, 1690, 650),
    }
    rounded_box(draw, boxes["input"], "#FFFFFF", "#9FB3C8", "状态输入", "F, D, Cr, N_n, h_n")
    rounded_box(draw, boxes["feature"], "#E8F5E9", "#66A06F", "模块一：物理派生特征", "F/D², Cr/D, log1p(N), h_n")
    rounded_box(draw, boxes["stress"], "#E3F2FD", "#4F83CC", "StressNet", "Transformer Encoder\n输出接触应力 p_n")
    rounded_box(draw, boxes["archard"], "#FFF3E0", "#D99030", "Archard 更新", "Δh_n = k · p_n · Δs_n\nh_{n+1}=h_n+Δh_n")
    rounded_box(draw, boxes["state"], "#F3E5F5", "#8E5AA8", "下一状态", "N_{n+1}, h_{n+1}\n达到 5 μm 输出寿命")
    arrow(draw, (345, 565), (420, 565))
    arrow(draw, (700, 565), (780, 565))
    arrow(draw, (1080, 565), (1150, 565))
    arrow(draw, (1410, 565), (1480, 565))

    # Feedback loop
    draw.line([(1588, 650), (1588, 735), (220, 735), (220, 650)], fill="#1F4E79", width=4)
    arrow(draw, (220, 735), (220, 650), color="#1F4E79", width=4)
    draw.text((620, 705), "闭环递推：上一时刻预测磨损参与下一时刻应力预测", font=get_font(22), fill="#1F4E79")

    # Loss block
    rounded_box(draw, (120, 805, 400, 925), "#FFFFFF", "#A0AEC0", "有限元真值", "p_true, h_true")
    rounded_box(draw, (535, 790, 1005, 940), "#FFFFFF", "#A0AEC0", "训练损失", "压力 MSE + 基础物理敏感性正则\n模块二：slow_abs 趋势约束")
    rounded_box(draw, (1135, 805, 1600, 925), "#FFFFFF", "#A0AEC0", "优化目标", "降低应力误差与闭环寿命误差\n提升跨工况泛化能力")
    arrow(draw, (400, 865), (535, 865), color="#596A7A")
    arrow(draw, (1005, 865), (1135, 865), color="#596A7A")
    draw.line([(930, 790), (930, 690), (930, 650)], fill="#596A7A", width=3)
    arrow(draw, (930, 690), (930, 650), color="#596A7A", width=3)

    img.save(path, quality=95)


def make_dataset_figure(summary: pd.DataFrame, split: pd.DataFrame, path: Path) -> None:
    if "D" in split.columns:
        df = split.copy()
    else:
        df = split.merge(summary[["file_name", "F", "D", "Cr"]], on="file_name", how="left")
    counts = df.groupby(["D", "split_role"]).size().unstack(fill_value=0).sort_index()
    img = Image.new("RGB", (1800, 920), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(36, True)
    label_font = get_font(22)
    small_font = get_font(19)
    draw.text((70, 45), "数据集划分与测试工况覆盖", font=title_font, fill="#102A43")
    draw.text((72, 95), "30 组仿真数据：26 组训练，4 组测试；四个 D 型号各保留一组测试工况", font=label_font, fill="#52616B")

    chart_x, chart_y = 110, 190
    bar_w, gap = 210, 80
    max_count = int(counts.max().max()) + 1
    axis_h = 470
    draw.line([(chart_x - 30, chart_y + axis_h), (chart_x + 4 * (bar_w + gap), chart_y + axis_h)], fill="#334E68", width=3)
    draw.line([(chart_x - 30, chart_y), (chart_x - 30, chart_y + axis_h)], fill="#334E68", width=3)
    for i, d in enumerate(counts.index):
        x = chart_x + i * (bar_w + gap)
        train = int(counts.loc[d].get("train", 0))
        test = int(counts.loc[d].get("test", 0))
        train_h = int(axis_h * train / max_count)
        test_h = int(axis_h * test / max_count)
        y0 = chart_y + axis_h - train_h
        draw.rounded_rectangle((x, y0, x + bar_w, chart_y + axis_h), radius=10, fill="#7BBF8E")
        draw.rounded_rectangle((x, y0 - test_h, x + bar_w, y0), radius=10, fill="#E76F51")
        draw.text((x + 55, chart_y + axis_h + 24), f"D={int(d)}", font=label_font, fill="#102A43")
        draw.text((x + 62, y0 + 12), f"训练 {train}", font=small_font, fill="white")
        draw.text((x + 70, y0 - test_h + 8), f"测试 {test}", font=small_font, fill="white")
    draw.rectangle((1080, 180, 1690, 690), outline="#D8DEE9", width=3, fill="#F8FAFC")
    draw.text((1115, 215), "固定测试集", font=get_font(28, True), fill="#1F4E79")
    test_df = df[df["split_role"].eq("test")].sort_values("D")
    y = 275
    for _, r in test_df.iterrows():
        draw.rounded_rectangle((1115, y, 1655, y + 82), radius=14, fill="white", outline="#C5D5E4", width=2)
        draw.text((1140, y + 13), f"{r['file_name'].replace('.csv', '')}", font=get_font(22, True), fill="#102A43")
        draw.text((1270, y + 13), f"F={r['F']:.0f}, D={r['D']:.0f}, Cr={r['Cr']:.3f}", font=small_font, fill="#334E68")
        draw.text((1270, y + 43), f"寿命={float(r['actual_life']):.1f} cycles", font=small_font, fill="#52616B")
        y += 100
    draw.rounded_rectangle((110, 735, 1660, 830), radius=18, fill="#EAF2F8", outline="#C5D5E4", width=2)
    note = "划分原则：测试集覆盖 D=8、13、16、22 四类尺寸，避免重复冲突工况；后续基准对比、模块筛选、消融实验和最终横向对比均使用同一固定划分。"
    for idx, line in enumerate(wrap_text(draw, note, label_font, 1480)):
        draw.text((145, 760 + idx * 31), line, font=label_font, fill="#243B53")
    img.save(path, quality=95)


def hex_to_rgb(value: str) -> tuple[int, int, int]:
    value = value.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def interp_color(stops: list[str], t: float) -> tuple[int, int, int]:
    t = max(0.0, min(1.0, t))
    if t >= 1.0:
        return hex_to_rgb(stops[-1])
    scaled = t * (len(stops) - 1)
    idx = int(math.floor(scaled))
    local = scaled - idx
    c0 = hex_to_rgb(stops[idx])
    c1 = hex_to_rgb(stops[idx + 1])
    return tuple(int(c0[i] + (c1[i] - c0[i]) * local) for i in range(3))


def make_per_run_error_heatmap(detail: pd.DataFrame, path: Path) -> None:
    run_order = ["Run4.csv", "Run11.csv", "Run25.csv", "Run28.csv"]
    model_order = ["Enhanced Transformer", "LSTM", "GRU", "Transformer", "1D-CNN", "FNN"]
    labels = [name.replace(".csv", "") for name in run_order]

    pivot = (
        detail.groupby(["variant", "file_name"], as_index=False)["life_rel_error"]
        .mean()
        .pivot(index="variant", columns="file_name", values="life_rel_error")
        .reindex(index=model_order, columns=run_order)
        * 100.0
    )

    img = Image.new("RGB", (1800, 980), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(38, True)
    subtitle_font = get_font(22)
    axis_font = get_font(23, True)
    cell_font = get_font(24, True)
    note_font = get_font(19)

    draw.text((90, 42), "不同测试工况寿命预测相对误差热力图", font=title_font, fill="#102A43")
    draw.text((92, 104), "列顺序按 Run 编号递增：Run4 -> Run11 -> Run25 -> Run28；数值为 5 个随机种子的平均相对误差", font=subtitle_font, fill="#52616B")

    left, top = 360, 240
    cell_w, cell_h = 260, 96
    row_label_w = 245
    palette = ["#FFF1C2", "#F8CF7A", "#F2A45B", "#E56F4A", "#B9574B"]
    values = [v for row in pivot.values.tolist() for v in row if pd.notna(v)]
    vmin, vmax = min(values), max(values)
    span = max(vmax - vmin, 1e-9)

    # Headers
    draw.rounded_rectangle((left - row_label_w, top - cell_h, left + cell_w * len(labels), top), radius=10, fill="#F4F7F8", outline="#D8DEE9", width=2)
    for j, label in enumerate(labels):
        x = left + j * cell_w
        draw.rectangle((x, top - cell_h, x + cell_w, top), fill="#F4F7F8", outline="#D8DEE9", width=2)
        tw, th = text_size(draw, label, axis_font)
        draw.text((x + (cell_w - tw) / 2, top - cell_h / 2 - th / 2), label, font=axis_font, fill="#243B53")

    for i, model in enumerate(model_order):
        y = top + i * cell_h
        draw.rectangle((left - row_label_w, y, left, y + cell_h), fill="#F8FAFC", outline="#D8DEE9", width=2)
        display_model = model.replace("Enhanced Transformer", "Enhanced\nTransformer")
        lines = display_model.split("\n")
        line_h = text_size(draw, "测试", axis_font)[1] + 5
        ly = y + (cell_h - line_h * len(lines)) / 2
        for line in lines:
            tw, th = text_size(draw, line, axis_font)
            draw.text((left - row_label_w + row_label_w - tw - 22, ly), line, font=axis_font, fill="#243B53")
            ly += line_h

        for j, run in enumerate(run_order):
            x = left + j * cell_w
            val = float(pivot.loc[model, run])
            t = (val - vmin) / span
            fill = interp_color(palette, t)
            draw.rectangle((x, y, x + cell_w, y + cell_h), fill=fill, outline="white", width=4)
            text_color = "#FFFFFF" if t > 0.58 else "#17324D"
            text = f"{val:.2f}%"
            tw, th = text_size(draw, text, cell_font)
            draw.text((x + (cell_w - tw) / 2, y + (cell_h - th) / 2), text, font=cell_font, fill=text_color)

    # Color legend
    legend_x, legend_y = left + cell_w * len(labels) + 80, top
    legend_w, legend_h = 58, cell_h * len(model_order)
    for k in range(legend_h):
        t = 1 - k / max(legend_h - 1, 1)
        color = interp_color(palette, t)
        draw.line([(legend_x, legend_y + k), (legend_x + legend_w, legend_y + k)], fill=color, width=1)
    draw.rectangle((legend_x, legend_y, legend_x + legend_w, legend_y + legend_h), outline="#C5D5E4", width=2)
    draw.text((legend_x - 6, legend_y - 34), "Error (%)", font=note_font, fill="#52616B")
    draw.text((legend_x + 76, legend_y - 4), f"{vmax:.1f}%", font=note_font, fill="#334E68")
    draw.text((legend_x + 76, legend_y + legend_h - 25), f"{vmin:.1f}%", font=note_font, fill="#334E68")

    note = "颜色采用黄-橙-红递增色带，颜色越偏橙红表示寿命相对误差越大。"
    draw.rounded_rectangle((left - row_label_w, 850, legend_x + 240, 925), radius=16, fill="#FFF9ED", outline="#E8CFA2", width=2)
    draw.text((left - row_label_w + 24, 872), note, font=subtitle_font, fill="#4B3428")

    img.save(path, quality=95)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).twips)))
    cell.width = Cm(width_cm)


def set_fixed_table_layout(table) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_run_font(run, font_name: str = BODY_FONT, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para_font(paragraph, font_name: str = BODY_FONT, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run, font_name, size, bold, color)


def add_body(doc: Document, text: str, first_line: bool = True) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(5)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.add_run(text)
    set_para_font(p, BODY_FONT, 10.5)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        set_run_font(run, HEAD_FONT, 15 if level == 1 else 12.5, True, ACCENT if level == 1 else "243B53")


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, BODY_FONT, 9.2, False, "52616B")


def add_figure(doc: Document, img_path: Path, caption: str, width_cm: float = 15.6) -> None:
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        run = p.add_run()
        run.add_picture(str(img_path), width=Cm(width_cm))
        add_caption(doc, caption)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[图位预留：{caption}。建议补充图片：{img_path.name}]")
        set_run_font(run, BODY_FONT, 10, True, "A61B1B")


def add_table(doc: Document, headers: list[str], rows: list[list[object]], caption: str, col_widths_cm: list[float] | None = None) -> None:
    global TABLE_IMG_COUNTER
    TABLE_IMG_COUNTER += 1
    img_path = FIG_DIR / f"table_{TABLE_IMG_COUNTER:02d}.png"
    make_table_image(headers, rows, caption, img_path, col_widths_cm)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Cm(15.5))
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, HEAD_FONT, 10.5, True, ACCENT)
    p2 = doc.add_paragraph()
    p2.paragraph_format.line_spacing = 1.25
    p2.paragraph_format.left_indent = Cm(0.45)
    p2.paragraph_format.right_indent = Cm(0.45)
    p2.paragraph_format.space_after = Pt(8)
    p2.add_run(body)
    set_para_font(p2, BODY_FONT, 9.8, False, "243B53")


def make_table_image(headers: list[str], rows: list[list[object]], caption: str, path: Path, col_widths_cm: list[float] | None = None) -> None:
    total_width = 2200
    margin_x = 36
    if col_widths_cm and sum(col_widths_cm) > 0:
        weights = [w / sum(col_widths_cm) for w in col_widths_cm]
    else:
        weights = [1 / len(headers)] * len(headers)
    widths = [int((total_width - 2 * margin_x) * w) for w in weights]
    widths[-1] += total_width - 2 * margin_x - sum(widths)

    tmp = Image.new("RGB", (total_width, 200), "white")
    draw = ImageDraw.Draw(tmp)
    header_font = get_font(28, True)
    body_font = get_font(25)
    pad_x = 18
    pad_y = 14

    def cell_lines(text: object, font: ImageFont.ImageFont, width: int) -> list[str]:
        return wrap_text(draw, str(text), font, max(40, width - 2 * pad_x))

    all_rows: list[list[list[str]]] = []
    header_lines = [cell_lines(h, header_font, widths[i]) for i, h in enumerate(headers)]
    all_rows.append(header_lines)
    for row in rows:
        all_rows.append([cell_lines(row[i], body_font, widths[i]) for i in range(len(headers))])

    row_heights: list[int] = []
    for idx, row_lines in enumerate(all_rows):
        font = header_font if idx == 0 else body_font
        _, line_h = text_size(draw, "测试", font)
        max_lines = max(len(lines) for lines in row_lines)
        row_heights.append(max(58, pad_y * 2 + max_lines * (line_h + 8)))

    caption_h = 68
    total_height = caption_h + sum(row_heights) + 2 * pad_y + 10
    img = Image.new("RGB", (total_width, total_height), "white")
    draw = ImageDraw.Draw(img)
    x0 = margin_x
    cap_font = get_font(26, True)
    cap_w, cap_h = text_size(draw, caption, cap_font)
    draw.text(((total_width - cap_w) // 2, 14), caption, font=cap_font, fill="#52616B")
    y = caption_h + pad_y

    for r_idx, row_lines in enumerate(all_rows):
        x = x0
        row_h = row_heights[r_idx]
        fill = f"#{ACCENT}" if r_idx == 0 else ("#F8FAFC" if r_idx % 2 == 0 else "#FFFFFF")
        text_color = "#FFFFFF" if r_idx == 0 else "#243B53"
        font = header_font if r_idx == 0 else body_font
        for c_idx, lines in enumerate(row_lines):
            w = widths[c_idx]
            draw.rectangle((x, y, x + w, y + row_h), fill=fill, outline=f"#{GRID}", width=2)
            _, line_h = text_size(draw, "测试", font)
            text_total_h = len(lines) * (line_h + 8) - 8
            ty = y + (row_h - text_total_h) // 2
            for line in lines:
                tw, _ = text_size(draw, line, font)
                if c_idx == 0 and r_idx != 0:
                    tx = x + pad_x
                else:
                    tx = x + (w - tw) // 2
                draw.text((tx, ty), line, font=font, fill=text_color)
                ty += line_h + 8
            x += w
        y += row_h

    img.save(path, quality=95)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        st = styles[style_name]
        st.font.name = HEAD_FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), HEAD_FONT)
        st.font.bold = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("寿命预测系统架构论文初稿")
    set_para_font(footer, BODY_FONT, 8.5, False, "7B8794")


def fmt_num(x: float, decimals: int = 1) -> str:
    return f"{float(x):.{decimals}f}"


def main() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)

    summary = pd.read_csv(DATA_SUMMARY)
    split = pd.read_csv(SPLIT_CSV)
    split_full = split.merge(summary[["file_name", "F", "D", "Cr"]], on="file_name", how="left")

    make_architecture_diagram(FIG_DIR / "Fig0_architecture_rebuilt.png")
    make_dataset_figure(summary, split_full, FIG_DIR / "Fig_data_split.png")

    image_map = {
        "Fig1_baseline_model_comparison.png": BEAUTY_DIR / "Fig1_baseline_model_comparison.png",
        "Fig2b_module1_refine_top2.png": BEAUTY_DIR / "Fig2b_module1_refine_top2.png",
        "Fig3b_module2_refine.png": BEAUTY_DIR / "Fig3b_module2_refine.png",
        "Fig4_ablation_mean_std.png": BEAUTY_DIR / "Fig4_ablation_mean_std.png",
        "Fig5_final_comparison.png": BEAUTY_DIR / "Fig5_final_comparison.png",
        "Fig6_final_per_run_error_heatmap.png": BEAUTY_DIR / "Fig6_final_per_run_error_heatmap.png",
    }
    for name, src in image_map.items():
        dst = FIG_DIR / name
        if src.exists():
            shutil.copyfile(src, dst)

    baseline = pd.read_csv(RESULT_DIR / "01_五模型基准对比" / "汇总_各模型平均指标.csv").sort_values("mean_life_abs_error")
    m1 = pd.read_csv(RESULT_DIR / "02_模块一_物理派生特征比较" / "复筛" / "汇总_各模型平均指标.csv").sort_values("mean_life_abs_error")
    m2 = pd.read_csv(RESULT_DIR / "03_模块二_趋势约束比较" / "复筛" / "汇总_各模型平均指标.csv").sort_values("mean_life_abs_error")
    ablation = pd.read_csv(RESULT_DIR / "04_最终消融实验" / "汇总_各模型平均指标.csv")
    final_cmp = pd.read_csv(RESULT_DIR / "05_最终横向对比" / "汇总_各模型平均指标.csv").sort_values("mean_life_abs_error")
    detail = pd.read_csv(RESULT_DIR / "05_最终横向对比" / "详细结果_各模型各测试集.csv")
    make_per_run_error_heatmap(detail, FIG_DIR / "Fig6_final_per_run_error_heatmap.png")

    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    r = title.add_run("寿命预测系统架构与物理增强模型研究")
    set_run_font(r, HEAD_FONT, 24, True, ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run("论文初稿（不含集成软件部分）")
    set_run_font(r, HEAD_FONT, 15, False, "52616B")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.line_spacing = 1.45
    meta_text = (
        "数据来源：完整 30 run 有限元仿真数据\n"
        "固定测试集：Run4 / Run28 / Run25 / Run11；训练集：其余 26 组 run\n"
        "统一设置：seq_len=12，闭环预测步长 cap=250，磨损阈值=5 μm，磨损系数=1.84e-10"
    )
    meta.add_run(meta_text)
    set_para_font(meta, BODY_FONT, 10.5, False, "334E68")

    doc.add_page_break()

    add_heading(doc, "写作说明", 1)
    add_callout(
        doc,
        "初稿定位",
        "本文档按照“数据集展示 - 架构介绍 - 物理信息增强 - 消融实验 - 整体效果展示”的逻辑撰写，适合作为论文中“寿命预测系统架构与模型验证”部分的初稿。后续若并入完整论文，可统一调整章节编号、图表编号和参考文献格式。",
    )

    add_heading(doc, "3.1 数据集构建与训练测试划分", 1)
    add_heading(doc, "3.1.1 仿真数据来源与变量说明", 2)
    add_body(
        doc,
        "本文寿命预测模型基于涂层关节轴承磨损有限元仿真数据构建。数据集共包含 30 组不同工况的仿真结果，每组数据记录了载荷 F、轴承特征尺寸 D、径向间隙 Cr、循环次数、磨损深度以及对应接触压力等信息。有限元仿真结果提供了模型训练所需的应力演化样本，同时根据磨损深度达到 5 μm 的时刻定义轴承寿命。",
    )
    add_body(
        doc,
        "从数据统计范围看，载荷 F 覆盖 500 N 至 9000 N，轴承尺寸 D 包含 8、13、16、22 四个水平，径向间隙 Cr 覆盖 0.008 至 0.05。各工况实际寿命范围约为 1.39×10^4 至 1.60×10^5 cycles，能够覆盖高载荷短寿命、低载荷长寿命以及中间过渡工况。由于不同工况的寿命跨度较大，模型不仅需要学习接触应力与磨损速率之间的关系，还需要在闭环递推过程中保持较好的稳定性。",
    )
    summary_rows = [
        ["仿真工况数量", f"{len(summary)} 组"],
        ["总采样行数", f"{int(summary['row_count'].sum())} 行"],
        ["载荷 F 范围", f"{summary['F'].min():.0f} - {summary['F'].max():.0f} N"],
        ["轴承尺寸 D 水平", ", ".join(str(int(x)) for x in sorted(summary["D"].unique()))],
        ["径向间隙 Cr 范围", f"{summary['Cr'].min():.3f} - {summary['Cr'].max():.3f}"],
        ["实际寿命范围", f"{summary['actual_life'].min():.1f} - {summary['actual_life'].max():.1f} cycles"],
        ["磨损阈值", "5 μm"],
    ]
    add_table(doc, ["统计项", "数值"], summary_rows, "表 3-1 数据集基本统计信息", [5.2, 8.5])

    add_heading(doc, "3.1.2 训练集与测试集划分", 2)
    add_body(
        doc,
        "为保证后续基准模型比较、增强模块筛选和消融实验具有一致的评价口径，本文采用固定训练测试划分。测试集选取 Run4、Run28、Run25 和 Run11 四组工况，其余 26 组工况作为训练集。该划分保证四个 D 型号各有一组测试样本，能够检验模型在不同轴承尺寸上的泛化能力；同时避免了部分工况中 F 与 Cr 完全重复导致的测试解释冲突。",
    )
    test_df = split_full[split_full["split_role"].eq("test")].sort_values("D")
    test_rows = []
    for _, r in test_df.iterrows():
        test_rows.append([
            r["file_name"].replace(".csv", ""),
            f"{r['F']:.0f}",
            f"{r['D']:.0f}",
            f"{r['Cr']:.3f}",
            fmt_num(r["actual_life"], 1),
            fmt_num(r["final_wear_um"], 3),
        ])
    add_table(
        doc,
        ["Run", "F/N", "D", "Cr", "实际寿命/cycles", "末端磨损/μm"],
        test_rows,
        "表 3-2 固定测试集参数",
        [2.0, 2.0, 1.5, 1.8, 3.2, 2.6],
    )
    add_figure(doc, FIG_DIR / "Fig_data_split.png", "图 3-1 数据集训练测试划分与测试工况覆盖", 15.6)

    add_heading(doc, "3.2 寿命预测系统总体架构", 1)
    add_heading(doc, "3.2.1 传统神经网络寿命预测模型", 2)
    add_body(
        doc,
        "传统神经网络寿命预测方法通常以载荷、结构尺寸和间隙等工况参数作为输入，直接回归预测寿命。这种方式具有建模简单、计算速度快的优点，但预测过程对磨损演化机理表达不足，模型输出与中间物理状态之间缺少可解释联系。当训练样本数量有限、测试工况与训练工况存在差异时，直接寿命回归模型容易受到数据分布影响，难以稳定反映磨损逐步累积直至失效的过程。",
    )
    add_heading(doc, "3.2.2 物理增强代理模型架构", 2)
    add_body(
        doc,
        "针对上述问题，本文采用物理增强代理模型构建寿命预测系统。模型不直接回归寿命，而是首先根据当前工况与磨损状态预测接触应力，再将预测应力代入 Archard 磨损公式计算磨损增量，并不断更新循环次数和磨损深度。当预测磨损深度达到 5 μm 阈值时，对应循环次数即为预测寿命。该结构把神经网络预测能力与磨损物理过程连接起来，使寿命预测结果能够通过中间应力和磨损曲线进行解释。",
    )
    add_figure(doc, FIG_DIR / "Fig0_architecture_rebuilt.png", "图 3-2 寿命预测系统总体架构与物理增强代理模型", 15.8)

    add_heading(doc, "3.2.3 闭环递推与寿命阈值判定", 2)
    add_body(
        doc,
        "在闭环预测过程中，系统以初始磨损深度 h0=0 作为起点。对于第 n 个递推步，模型根据 F、D、Cr、当前循环次数 Nn 和当前磨损深度 hn 预测接触应力 pn。随后根据 Archard 磨损关系计算磨损增量：Δhn = k · pn · Δsn，其中 k 为磨损系数，Δsn 表示该步对应的滑移距离。本文统一采用磨损系数 1.84e-10，并将闭环递推的实际循环步长上限设为 250 cycles，以避免不同工况原始采样步长差异对寿命预测造成额外影响。",
    )
    add_body(
        doc,
        "递推得到的下一步状态为 hn+1 = hn + Δhn，Nn+1 = Nn + ΔN。若 hn+1 尚未达到 5 μm，则将新状态重新输入模型进行下一步预测；若 hn+1 超过阈值，则在当前步内按线性插值确定达到阈值的循环次数。通过这种方式，最终寿命不是单个神经网络输出值，而是由应力预测和磨损积分共同决定。",
    )

    add_heading(doc, "3.3 物理信息增强方法", 1)
    add_body(
        doc,
        "在总体架构基础上，本文围绕 Transformer 模型设计两个增强模块。第一个模块从输入端引入物理派生特征，使模型更容易识别载荷、尺寸与间隙之间的组合影响；第二个模块从训练目标端引入趋势约束，使应力预测在同一工况的时间序列上符合磨损过程的总体演化规律。两个模块分别对应特征表达增强和损失函数增强，最终共同形成 Enhanced Transformer。",
    )
    add_heading(doc, "3.3.1 模块一：物理派生特征构造", 2)
    add_body(
        doc,
        "原始输入特征包括 F、D、Cr、actual_cycle 和 wear_depth 五个变量。虽然这些变量包含寿命预测所需的基本信息，但神经网络需要自行学习载荷、尺寸和间隙之间的非线性组合关系。考虑到接触应力与载荷分布面积、结构尺度及间隙之间存在明显物理联系，本文构造 F/D² 和 Cr/D 两个静态派生特征，以显式表达载荷密度和相对间隙效应。",
    )
    add_body(
        doc,
        "此外，由于循环次数跨度较大，直接输入 actual_cycle 可能导致模型对长寿命工况的尺度变化过于敏感。本文进一步引入 log1p(actual_cycle)，用对数形式压缩循环次数尺度，使模型更关注磨损演化阶段而非绝对循环数值。模块一筛选中比较了原始五维输入、静态派生特征、log_cycle 替换以及 log_cycle 保留等方案。复筛结果显示，M1_R3_log_cycle_keep 和 M1_R2_log_cycle_replace 均明显优于原始 Transformer 输入，其中 M1_R3 单独表现略优，但后续与趋势约束联合筛选时，M1_R2 与模块二组合后的综合效果更好，因此最终增强模型采用 M1_R2_log_cycle_replace。",
    )
    m1_rows = []
    for _, r in m1.iterrows():
        m1_rows.append([
            r["variant"],
            fmt_num(r["mean_life_abs_error"], 1),
            fmt_num(r["std_life_abs_error"], 1),
            fmt_num(r["mean_pressure_mae"], 3),
            fmt_num(r["mean_wear_mae_um"], 4),
        ])
    add_table(doc, ["模块一方案", "平均寿命误差", "标准差", "压力 MAE", "磨损 MAE/μm"], m1_rows, "表 3-3 模块一复筛结果", [5.0, 2.7, 2.0, 2.0, 2.3])
    add_figure(doc, FIG_DIR / "Fig2b_module1_refine_top2.png", "图 3-3 模块一物理派生特征复筛对比", 15.0)

    add_heading(doc, "3.3.2 模块二：应力趋势约束损失", 2)
    add_body(
        doc,
        "在磨损递推过程中，接触应力预测的局部波动会被持续积分放大，最终影响寿命阈值达到时刻。因此，除了单点压力误差外，还需要约束同一工况时间序列上的应力演化形态。本文设计的模块二不是要求应力严格逐点单调下降，而是采用温和的 slow_abs 趋势约束，鼓励应力下降幅度保持平缓，避免相邻时刻之间出现不合理的剧烈变化。",
    )
    add_body(
        doc,
        "具体而言，对于同一 run 内连续三个状态的预测应力 s0、s1、s2，定义相邻下降量 drop0=s0-s1，drop1=s1-s2。当后一段下降量明显大于前一段时，通过 ReLU(drop1-drop0) 产生惩罚。该约束能够抑制闭环预测中不稳定的应力突变，同时保留有限元数据中可能存在的局部波动。模块二比较了无额外趋势约束、slow_abs 以及 temporal_mono_plus_slow_abs 等方案，复筛结果表明 slow_abs_0p01 表现最优。",
    )
    m2_rows = []
    for _, r in m2.head(3).iterrows():
        m2_rows.append([
            r["variant"],
            fmt_num(r["mean_life_abs_error"], 1),
            fmt_num(r["std_life_abs_error"], 1),
            fmt_num(r["mean_pressure_mae"], 3),
            fmt_num(r["mean_wear_mae_um"], 4),
        ])
    add_table(doc, ["模块二组合方案", "平均寿命误差", "标准差", "压力 MAE", "磨损 MAE/μm"], m2_rows, "表 3-4 模块二复筛 Top3 结果", [6.0, 2.5, 1.8, 1.8, 2.0])
    add_figure(doc, FIG_DIR / "Fig3b_module2_refine.png", "图 3-4 模块二趋势约束复筛对比", 15.0)

    add_heading(doc, "3.4 消融实验设计与结果分析", 1)
    add_body(
        doc,
        "为验证两个增强模块的独立贡献，本文设计四组消融实验：T0_baseline 表示原始 Transformer；T1_module1_only 仅加入模块一物理派生特征；T2_module2_only 仅加入模块二趋势约束；T3_module1_plus_module2 同时加入两个模块。所有消融实验均采用相同固定测试集、相同闭环递推设置和 5 个随机种子重复训练。",
    )
    order = ["T0_baseline", "T1_module1_only", "T2_module2_only", "T3_module1_plus_module2"]
    ablation = ablation.set_index("variant").loc[order].reset_index()
    base_error = float(ablation.iloc[0]["mean_life_abs_error"])
    ab_rows = []
    for _, r in ablation.iterrows():
        delta = float(r["mean_life_abs_error"]) - base_error
        ab_rows.append([
            r["variant"],
            fmt_num(r["mean_life_abs_error"], 1),
            fmt_num(r["std_life_abs_error"], 1),
            f"{delta:+.1f}",
            fmt_num(r["mean_pressure_mae"], 3),
            fmt_num(r["mean_wear_mae_um"], 4),
        ])
    add_table(doc, ["消融方案", "平均寿命误差", "标准差", "相对 T0 变化", "压力 MAE", "磨损 MAE/μm"], ab_rows, "表 3-5 最终消融实验结果", [4.2, 2.4, 1.8, 2.2, 1.8, 2.0])
    add_body(
        doc,
        "消融结果表明，模块一单独使用时，平均寿命绝对误差由 1295.2 cycles 降低至 988.2 cycles，说明物理派生特征能够改善 Transformer 对不同工况参数组合的表达能力。模块二单独使用时，误差进一步降至 810.0 cycles，表明趋势约束对闭环寿命递推稳定性具有更直接的改善作用。当两个模块联合使用时，平均寿命误差降至 607.3 cycles，相比原始 Transformer 减少 687.9 cycles，达到四组消融方案中的最优结果。",
    )
    add_body(
        doc,
        "值得注意的是，T2_module2_only 的磨损曲线 MAE 最低，而 T3_module1_plus_module2 的寿命误差最低。这说明磨损曲线局部误差与阈值寿命误差并非完全等价：寿命预测更依赖闭环递推全过程中误差累积的方向和稳定性。模块一与模块二联合后，模型在特征表达和趋势稳定性两方面同时受益，因此最终寿命预测效果最佳。",
    )
    add_figure(doc, FIG_DIR / "Fig4_ablation_mean_std.png", "图 3-5 最终四格消融实验结果", 15.0)

    add_heading(doc, "3.5 整体预测效果展示", 1)
    add_heading(doc, "3.5.1 与基准模型的横向对比", 2)
    add_body(
        doc,
        "为了评估最终增强模型的综合性能，本文将 Enhanced Transformer 与 FNN、GRU、LSTM、1D-CNN 和原始 Transformer 进行横向比较。所有模型均在相同训练测试划分和统一闭环递推口径下评价，主要指标包括平均寿命绝对误差、接触压力 MAE 和磨损曲线 MAE。",
    )
    final_rows = []
    for _, r in final_cmp.iterrows():
        final_rows.append([
            r["variant"],
            fmt_num(r["mean_life_abs_error"], 1),
            fmt_num(r["std_life_abs_error"], 1),
            fmt_num(r["mean_pressure_mae"], 3),
            fmt_num(r["mean_wear_mae_um"], 4),
        ])
    add_table(doc, ["模型", "平均寿命误差", "标准差", "压力 MAE", "磨损 MAE/μm"], final_rows, "表 3-6 最终横向对比结果", [4.2, 2.5, 1.9, 2.0, 2.0])
    add_body(
        doc,
        "结果显示，Enhanced Transformer 的平均寿命绝对误差为 607.3 cycles，在所有比较模型中最低。最强原始基准为 LSTM，其平均寿命误差为 892.8 cycles；相比之下，增强模型仍降低约 285.5 cycles。原始 Transformer 的误差为 1295.2 cycles，说明本文最终性能提升并非来自 Transformer 基础结构本身，而主要来自物理派生特征和趋势约束两个增强模块的联合作用。",
    )
    add_figure(doc, FIG_DIR / "Fig5_final_comparison.png", "图 3-6 Enhanced Transformer 与基准模型横向对比", 15.0)

    add_heading(doc, "3.5.2 不同测试工况下的误差分布", 2)
    per_run = (
        detail.groupby(["variant", "file_name"])
        .agg(
            mean_life_abs_error=("life_abs_error", "mean"),
            mean_rel_error=("life_rel_error", "mean"),
            mean_predicted_life=("predicted_life", "mean"),
            true_life=("true_life", "first"),
        )
        .reset_index()
    )
    enh_rows = []
    run_order = ["Run4.csv", "Run11.csv", "Run25.csv", "Run28.csv"]
    per_run_enh = per_run[per_run["variant"].eq("Enhanced Transformer")].copy()
    per_run_enh["run_order"] = per_run_enh["file_name"].map({name: i for i, name in enumerate(run_order)})
    for _, r in per_run_enh.sort_values("run_order").iterrows():
        enh_rows.append([
            r["file_name"].replace(".csv", ""),
            fmt_num(r["true_life"], 1),
            fmt_num(r["mean_predicted_life"], 1),
            fmt_num(r["mean_life_abs_error"], 1),
            f"{100 * float(r['mean_rel_error']):.2f}%",
        ])
    add_table(doc, ["测试工况", "真实寿命", "平均预测寿命", "平均绝对误差", "平均相对误差"], enh_rows, "表 3-7 Enhanced Transformer 在各测试工况上的预测表现", [2.2, 2.8, 3.0, 2.8, 2.4])
    add_body(
        doc,
        "从单独测试工况看，Enhanced Transformer 并非在每一个 run 上都取得绝对最小误差，但在固定测试集上的平均寿命误差最低。单工况误差受到工况边界、寿命长短、有限元采样步长以及闭环误差累积方向的共同影响，因此更适合用整体平均指标评价模型泛化能力。相对误差热力图进一步显示，增强模型在长寿命工况 Run4 和 Run28 上的误差较低，这对实际寿命预测具有重要意义。",
    )
    add_figure(doc, FIG_DIR / "Fig6_final_per_run_error_heatmap.png", "图 3-7 各测试工况寿命预测相对误差热力图", 15.2)

    add_heading(doc, "3.5.3 综合分析", 2)
    add_body(
        doc,
        "综合数据集划分、模块筛选、消融实验和最终横向对比可以看出，本文寿命预测系统的关键并不是单纯替换神经网络结构，而是将神经网络应力预测与磨损物理递推过程结合起来。在该框架下，Transformer 负责学习时序状态到接触应力的映射，Archard 公式负责将应力转化为磨损增量，闭环递推负责将局部应力预测累积为寿命结果。两个增强模块则分别从输入表达和训练目标两侧补充物理先验，使模型在有限样本条件下获得更稳定的泛化能力。",
    )
    add_body(
        doc,
        "因此，Enhanced Transformer 可以作为本文寿命预测系统的核心模型。与传统直接寿命回归方法相比，该方法保留了磨损演化路径；与原始 Transformer 相比，该方法显式引入物理派生特征并约束应力趋势；与 LSTM 等强基准相比，该方法在固定测试集上取得了更低的平均寿命误差。后续研究可进一步扩展测试工况数量，并结合真实试验数据对有限元代理模型进行修正，以提高模型在实际服役环境中的可靠性。",
    )

    add_heading(doc, "本章小结", 1)
    add_body(
        doc,
        "本章围绕涂层关节轴承磨损寿命预测问题，建立了不含集成软件部分的寿命预测系统架构初稿。首先介绍了完整 30 run 数据集及固定训练测试划分；随后提出基于 StressNet、Archard 磨损更新和闭环寿命阈值判定的物理增强代理模型；进一步设计了物理派生特征构造和应力趋势约束两个增强模块；最后通过消融实验和横向对比证明，Enhanced Transformer 在固定测试集上的平均寿命绝对误差最低，能够支撑本文“物理特征构造 + 趋势约束增强 Transformer”的主要结论。",
    )

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
